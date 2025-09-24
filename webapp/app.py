# webapp/app.py
from flask import Flask, render_template, request, send_from_directory, url_for, redirect
from pathlib import Path
import tempfile, time

from crewai import Agent, Task, Crew, Process, LLM

try:
    import docx
    from docx import Document
except ImportError:
    docx = None

app = Flask(__name__)
OUTPUT_DIR = Path(__file__).parent / "output"
OUTPUT_DIR.mkdir(exist_ok=True)

# ---------- helpers ----------
def load_minutes_from_path(p: Path) -> str:
    if not p.exists():
        raise FileNotFoundError(f"File not found: {p}")
    if p.suffix.lower() == ".txt":
        return p.read_text(encoding="utf-8", errors="ignore")
    if p.suffix.lower() == ".docx":
        if docx is None:
            raise RuntimeError("Install python-docx to read .docx files (pip install python-docx)")
        d = docx.Document(str(p))
        return "\n".join(par.text for par in d.paragraphs)
    raise ValueError("Unsupported file type. Please upload .txt or .docx")

def summarize_minutes(text: str) -> str:
    llm = LLM(model="ollama/llama3.2:1b", base_url="http://localhost:11434")

    summariser = Agent(
        name="Minutes Summariser",
        role="Summarises meeting minutes",
        goal="Return only key decisions, action items (with owners if present), and important discussion points.",
        backstory=(
            "You are concise and factual. Extract only what is present in the minutes. "
            "Never invent names or dates."
        ),
        llm=llm,
        verbose=False,
    )

    prompt = f"""
Summarise the following meeting minutes. Be concise and factual. Do not invent names or dates.
Return output in this format:

# Key Decisions Made:
- ...

# Action Items:
- ...

# Important Discussion Points:
- ...

Minutes:
\"\"\"{text[:20000]}\"\"\"  # truncated for safety
"""

    task = Task(
        description=prompt,
        agent=summariser,
        expected_output=(
            "A structured summary with three sections: Key Decisions Made, Action Items, "
            "and Important Discussion Points. Use headings and bullet points."
        ),
    )

    crew = Crew(agents=[summariser], tasks=[task], process=Process.sequential)
    result = crew.kickoff()
    return str(result)

def save_summary_txt(summary: str, original_name: str) -> Path:
    ts = time.strftime("%Y%m%d-%H%M%S")
    out_txt = OUTPUT_DIR / f"{Path(original_name).stem}_summary_{ts}.txt"
    out_txt.write_text(summary, encoding="utf-8")
    return out_txt

def save_summary_docx(summary: str, original_name: str) -> Path:
    if docx is None:
        raise RuntimeError("Install python-docx to write .docx files (pip install python-docx)")
    ts = time.strftime("%Y%m%d-%H%M%S")
    out_docx = OUTPUT_DIR / f"{Path(original_name).stem}_summary_{ts}.docx"

    doc = Document()
    doc.add_heading("Meeting Summary", level=1)
    for line in summary.splitlines():
        s = line.strip()
        if not s:
            continue
        if s.startswith("# "):
            doc.add_heading(s[2:].strip(), level=2)
        elif s.startswith("## "):
            doc.add_heading(s[3:].strip(), level=3)
        elif s.startswith("- "):
            doc.add_paragraph(s[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(s)
    doc.save(out_docx)
    return out_docx

# ---------- routes ----------
@app.route("/", methods=["GET"])
def index():
    return render_template("index.html", summary=None, download_docx_url=None)

@app.route("/summarise", methods=["POST"])
def summarise():
    f = request.files.get("file")
    if not f or f.filename == "":
        return render_template("index.html", summary="Please choose a .txt or .docx file.", download_docx_url=None)

    suffix = Path(f.filename).suffix.lower()
    if suffix not in {".txt", ".docx"}:
        return render_template("index.html", summary="Unsupported file type. Please upload .txt or .docx.", download_docx_url=None)

    # save upload to a temp file and read it
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        f.save(tmp.name)
        minutes_text = load_minutes_from_path(Path(tmp.name))

    # run summary
    summary = summarize_minutes(minutes_text)

    # save copies (txt + docx) to /output
    _ = save_summary_txt(summary, f.filename)
    docx_path = save_summary_docx(summary, f.filename)
    download_docx_url = url_for("download_file", filename=docx_path.name)

    return render_template("index.html", summary=summary, download_docx_url=download_docx_url)

@app.route("/download/<path:filename>", methods=["GET"])
def download_file(filename):
    return send_from_directory(OUTPUT_DIR, filename, as_attachment=True)

@app.route("/clear", methods=["GET"])
def clear_summary():
    return redirect(url_for("index"))

if __name__ == "__main__":
    app.run(debug=True, host="127.0.0.1", port=8000)