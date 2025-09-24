# Minimal Minutes by TamasFried
# -----------------------------
# This script is a basic meeting minutes summariser that captures key points from a meeting transcript.
# It uses:
#     - CrewAI (for agent/task assignment)
#     - LangChain (to connect to a local LLM)
#     - Ollama model "llama3.2:1b" (small and fast for local testing)
#
# Workflow:
#     1. Load the meeting minutes as .txt or .docx file.
#     2. Pass the text to a CrewAI agent (Llama in this case).
#     3. The Agent produces a structured summary of the meeting minutes, including:
#         - Key decisions made
#         - Action items assigned (with responsible persons)
#         - Important discussion points
#     4. Print the structured summary.
#
# Running this basic script:
#     python minimal_minutes.py minutes.txt
#     OR
#     python minimal_minutes.py minutes.docx

# Basic Imports
import sys                  # For command line arguments
from pathlib import Path    # For file path handling

# External Libraries
from crewai import Agent, Task, Crew, Process
from langchain_ollama import ChatOllama
try:
    import docx # type: ignore # For reading .docx files
except ImportError:
    docx = None # only needed if .docx files are used

# ----------------------
# File Loading Function
# ----------------------
def load_minutes(path_str: str) -> str:
    # Load text from a .txt or .docx file
    # Returns the content as a string.
    p = Path(path_str)
    if not p.exists():
        raise FileNotFoundError(f"File not found: {p}")
    
    # Handling plain text files
    if p.suffix.lower() == '.txt':
        return p.read_text(encoding='utf-8', errors='ignore')
    
    # Handling Word docs
    if p.suffix.lower() == ".docx":
        if docx is None:
            raise RuntimeError("python-docx is required for .docx files. Install with: pip install python-docx")
        d = docx.Document(str(p))   # create a Document object
        return "\n".join(par.text for par in d.paragraphs)  # join all paragraphs with newlines
    
    # If unsupported file type
    raise ValueError("Unsupported file type. Please provide a .txt or .docx file.")

# ----------------------
# File Saving Function
# ----------------------
from docx import Document # type: ignore
from pathlib import Path

def save_summary_to_docx(summary: str, input_file: str):
    """
    Save the model's summary output to a Word document inside an 'output' folder.
    Filename will be <inputname>_summary.docx
    """
    input_path = Path(input_file)

    # Make sure output folder exists
    output_dir = input_path.parent / "output"
    output_dir.mkdir(exist_ok=True)

    out_path = output_dir / f"{input_path.stem}_summary.docx"

    doc = Document()
    doc.add_heading("Meeting Summary", level=1)

    # Split summary into lines and format them nicely
    for line in summary.splitlines():
        if line.strip().startswith("# "):  # Heading
            doc.add_heading(line.strip("# ").strip(), level=2)
        elif line.strip().startswith("- "):  # Bullet point
            doc.add_paragraph(line.strip("- ").strip(), style="List Bullet")
        elif line.strip():
            doc.add_paragraph(line.strip())

    doc.save(out_path)
    print(f"✅ Summary saved to {out_path}")


# ----------------------
# Main Function
# ----------------------
def main():
    # Runs the summarisation process.

    # Ensure file path is provided
    if len(sys.argv) != 2:
        print("Usage: python minimal_minutes.py <minutes.txt or minutes.docx>")
        sys.exit(1)

    # Load the minutes text
    minutes_text = load_minutes(sys.argv[1])

    # Initialize the LLM (Ollama with Llama 3.2 1B model)
    llm = ChatOllama(
    model="ollama/llama3.2:1b"
)

    # Define the task for the agent
    summariser = Agent(
        name = "Minutes Summariser",
        role = "You are an expert meeting minutes summariser.",
        goal = "Return only key points from the meeting minutes, including decisions, action items with responsible persons, and important discussion points as well as generating a title for the meeting.",
        backstory="You are an assistant who specialises in summarising meeting notes. "
              "You extract only the key facts, decisions, and actions. "
              "You never invent details not present in the text.",
        llm = llm,
        verbose = False,
    )

    # Define the task/prompt for the agent
    prompt = f"""
Summarise the following minutes.
Do not invent any information such as names or dates.
Return the summary in the following format:

# Meeting Title: <Generate a concise title for the meeting>
# Key Decisions Made:
- Decision 1
- Decision 2 or more
# Action Items:
- Action Item 1 (Responsible Person)
- Action Item 2 (Responsible Person) or more
# Important Discussion Points:
- Point 1
- Point 2 or more
Here are the minutes:
\"\"\"
{minutes_text}\"\"\"
"""
    
    task = Task(
        description = prompt, 
        agent = summariser,
        expected_output = "A structured summary of the meeting minutes containing Key Decisions, "
                    "Action Items (with owners if mentioned), and Discussion Points. "
                    "The format must match the example given in the prompt."
        )

    # Create a Crew and run the task
    crew = Crew(agents=[summariser], tasks=[task], process=Process.sequential)
    result = crew.kickoff()

    # Print the structured summary
    print(result)

    # Save the summary to a Word document
    save_summary_to_docx(str(result), "meeting_summary.docx")

if __name__ == "__main__":
    main()